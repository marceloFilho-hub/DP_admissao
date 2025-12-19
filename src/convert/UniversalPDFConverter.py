import os
from pathlib import Path
from PIL import Image
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
import csv
from docx import Document


# =====================================================
# CONFIGURAÇÕES DE DIRETÓRIO
# =====================================================
BASE_DIR = Path(__file__).resolve().parents[2]

INPUT_DIR = BASE_DIR / "src" / "convert" / "arq_descompactado"

SUPPORTED_IMAGE_EXT = ["png", "jpg", "jpeg", "bmp", "tiff"]
SUPPORTED_TEXT_EXT = ["txt"]
SUPPORTED_CSV_EXT = ["csv"]
SUPPORTED_DOCX_EXT = ["docx"]


class UniversalPDFConverter:

    def _generate_output_path(self, input_path: str) -> Path:
        input_path = Path(input_path)
        return input_path.with_suffix(".pdf")  # => mantém mesma pasta

    def convert(self, filepath: str) -> Path:
        ext = Path(filepath).suffix.lower().replace(".", "")

        if ext == "pdf":
            return Path(filepath)
        
        if ext in SUPPORTED_DOCX_EXT:
            return self._convert_docx_to_pdf(filepath)


        if ext in SUPPORTED_IMAGE_EXT:
            return self._convert_image_to_pdf(filepath)

        if ext in SUPPORTED_TEXT_EXT:
            return self._convert_txt_to_pdf(filepath)

        if ext in SUPPORTED_CSV_EXT:
            return self._convert_csv_to_pdf(filepath)

        raise ValueError(f"Formato não suportado: {ext}")

    # =====================================================
    # 1. IMAGEM → PDF
    # =====================================================
    def _convert_image_to_pdf(self, input_path):
        pdf_path = self._generate_output_path(input_path)
        img = Image.open(input_path).convert("RGB")
        img.save(pdf_path)
        return pdf_path

    # =====================================================
    # 2. TXT → PDF
    # =====================================================
    def _convert_txt_to_pdf(self, input_path):
        pdf_path = self._generate_output_path(input_path)

        with open(input_path, "r", encoding="utf-8") as f:
            lines = f.readlines()

        c = canvas.Canvas(str(pdf_path), pagesize=A4)
        width, height = A4
        y = height - 40

        for line in lines:
            if y < 40:
                c.showPage()
                y = height - 40
            c.drawString(40, y, line.strip())
            y -= 20

        c.save()
        return pdf_path

    # =====================================================
    # 3. CSV → PDF
    # =====================================================
    def _convert_csv_to_pdf(self, input_path):
        pdf_path = self._generate_output_path(input_path)

        c = canvas.Canvas(str(pdf_path), pagesize=A4)
        width, height = A4
        y = height - 40

        with open(input_path, "r", encoding="utf-8") as f:
            reader = csv.reader(f)
            for row in reader:
                linha = " | ".join(row)
                if y < 40:
                    c.showPage()
                    y = height - 40
                c.drawString(40, y, linha)
                y -= 20

        c.save()
        return pdf_path

    # =====================================================
    # 4. DOCX → PDF
    # =====================================================
    def _convert_docx_to_pdf(self, input_path):
        pdf_path = self._generate_output_path(input_path)

        doc = Document(input_path)

        c = canvas.Canvas(str(pdf_path), pagesize=A4)
        width, height = A4
        y = height - 50

        def new_page():
            nonlocal y
            c.showPage()
            y = height - 50

        # Processa parágrafos
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()

            if text:
                if y < 50:
                    new_page()
                c.drawString(40, y, text)
                y -= 18

        # Processa tabelas
        for table in doc.tables:
            for row in table.rows:
                linha = " | ".join(
                    cell.text.strip().replace("\n", " ")
                    for cell in row.cells
                )

                if y < 50:
                    new_page()

                c.drawString(40, y, linha)
                y -= 18

        c.save()
        return pdf_path


# =====================================================
# PROCESSAMENTO RECURSIVO EM TODAS AS SUBPASTAS
# =====================================================
if __name__ == "__main__":

    converter = UniversalPDFConverter()

    print(f"[INFO] Escaneando pastas em: {INPUT_DIR}\n")

    # percorre TODAS as pastas recursivamente
    for root, dirs, files in os.walk(INPUT_DIR):

        for file in files:
            filepath = Path(root) / file
            ext = filepath.suffix.lower().replace(".", "")

            # ignora PDFs (já são finais)
            if ext == "pdf":
                continue

            try:
                print(f"[PROCESSANDO] {filepath}")
                pdf_path = converter.convert(str(filepath))

                # remove o arquivo original
                os.remove(filepath)

                print(f"[OK] Gerado: {pdf_path}")

            except Exception as e:
                print(f"[ERRO] Falha ao converter {filepath} → {e}")

    print("\n[TUDO CONCLUÍDO]")






